from flask import Flask, send_file
from docx import Document
import os

app = Flask(__name__)


@app.route('/')
def generate_document():
    # Create a new Word Document
    doc = Document()

    # Add a title
    doc.add_heading("Understanding Random Character Selection with Seed in Python", level=1)

    # Add an introduction
    doc.add_paragraph(
        "This document provides a detailed explanation of how the 'random.seed()' function "
        "works in Python and how it helps in consistently selecting a random character from "
        "a string based on a seed value. The associated Python code is included for reference."
    )

    # Add the Python code as a code block
    code = """
import random

def random_letter(ex_str, seed):
    # Set the seed for reproducible results
    random.seed(seed)

    # Select a random character from the string
    return random.choice(ex_str)

# Taking input from the user for the string and seed
input_string = input("Enter a string: ")
input_seed = int(input("Enter a seed value (integer): "))

# Calling the function and printing the result
result = random_letter(input_string, input_seed)
print(f"Random character from '{input_string}' with seed {input_seed}: {result}")
"""
    doc.add_heading("Python Code", level=2)
    doc.add_paragraph(code)

    # Add an explanation section
    doc.add_heading("Explanation", level=2)
    doc.add_paragraph(
        "1. The 'random.seed()' function initializes the random number generator with the "
        "given seed value. This ensures reproducibility, meaning the same random sequence "
        "will be generated for the same seed."
    )
    doc.add_paragraph(
        "2. The 'random.choice()' function selects a random element from the given string. "
        "When combined with a fixed seed, it will always select the same character for the "
        "same input string and seed value."
    )
    doc.add_paragraph(
        "3. The program takes input for a string and a seed value from the user, applies the "
        "random selection process, and outputs the selected character."
    )
    doc.add_paragraph(
        "For example:\n"
        "  - Input string: 'Python Rocks!'\n"
        "  - Seed value: 1\n"
        "  - Output: 't'"
    )

    # Save the document
    file_path = "Random_Character_Selection_Explanation.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
